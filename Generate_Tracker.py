#!/usr/bin/env python
# -*- coding: ascii -*-
"""
Build a tracker DOCX with:
 - Page 1: vehicle/BMS summary + reset counts/results (read from tracker_summary.csv).
 - Subsequent pages: one per test case, showing ASCII summary, graph image, and PASS/FAIL.

Relies on outputs produced by the test scripts (summary/result/graph) and file_name.json
to map filenames. Use python-docx to render the document.
"""

import os
import sys
import json
import csv
from typing import Dict, List, Tuple
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.image.image import Image

# ------------------------------------------------------------
# CONFIG (mirrors LAUNCHER.py)
# ------------------------------------------------------------
TEST_CASES = [
    "SoC BEHAVIOR", "SHUTDOWN PROCESS", "PRECHARGE PROCESS",
    "BMS STATE TRANSITION", "CELL TEMP IMBALANCE", "BMS PCB TEMP",
    "ANY BMS ERROR", "FLAG FULL CHARGE DISABLE", "DCLI / DCLO MAP",
    "EQUIVALENT CYCLE COUNT", "BMS BALANCING",
    "PRIMARY VS SECONDARY LATCH", "MCU OBC ERROR",
    "AuxCharge_with_Vehicle_state_change", "SoC vs VOLTAGE SUMMARY",
    "CAPACITY CHECK", "BMS CURRENT IN READY MODE", "DRIVE_CHARGE Max Min Avg CURRENT"
]

# Skip summary JSON for these tests (case-insensitive); still include graph/result.
SKIP_SUMMARY_FOR = {
    "shutdown process",
    "precharge process",
    "bms state transition",
    "any bms error",
    "mcu obc error",
    "soc vs voltage summary",
    "cell temp imbalance",
    "bms pcb temp",
}

SCRIPT_BY_ROW: Dict[int, str] = {
    0: "SoC_behavior.py",
    1: "Shutdown_Process.py",
    2: "Precharge_Process.py",
    3: "BMS_State_transition.py",
    4: "Cell_Temp_Imbalance.py",
    5: "BMS_PCB_Temp.py",
    6: "Any_BMS_Error.py",
    7: "Flag_Full_Charge_Disable.py",
    8: "DCLI_DCLO_Map.py",
    9: "Equivalent_cycle_count.py",
    10: "BMS_Balancing.py",
    11: "Primary_vs_Secondary_Latch.py",
    12: "MCU_OBC_Error.py",
    13: "AuxCharge_with_Vehicle_state_change.py",
    14: "SoC_vs_Voltage_Summary.py",
    15: "Capacity_check.py",
    16: "BMS_Current_in_Ready_Mode.py",
    17: "DRIVE_CHARGE_Max_Min_Avg_CURRENT.py",
}

DEFAULT_TESTS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TRC TEST CASES")
FILE_MAP_NAME = "file_name.json"
META_CSV = "tracker_summary.csv"
OUTPUT_DOCX = "tracker_summary.docx"
META_JSON_ENV = "META_JSON"
SELECTED_FILE_ENV = "SELECTED_FILE_NAME"
MAX_GRAPH_WIDTH_IN = 6.5
MAX_GRAPH_HEIGHT_IN = 4.5


def _default_output_names(script_name: str) -> Dict[str, str]:
    base = os.path.splitext(script_name)[0]
    return {
        "result": f"{base}_results.json",
        "summary": f"{base}_summary.json",
        "graph": f"{base}_plot.png",
    }


# ------------------------------------------------------------
# UTILITIES
# ------------------------------------------------------------
def load_meta(csv_path: str) -> Dict[str, str]:
    """Load vehicle/BMS/reset info into a dict. Env META_JSON takes priority over CSV."""
    meta_env = os.environ.get(META_JSON_ENV)
    if meta_env:
        try:
            parsed = json.loads(meta_env)
            if isinstance(parsed, dict):
                return {str(k): str(v) for k, v in parsed.items()}
        except Exception:
            pass

    meta: Dict[str, str] = {}
    if csv_path and os.path.exists(csv_path):
        try:
            with open(csv_path, newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    if len(row) >= 2:
                        meta[row[0].strip()] = row[1].strip()
        except Exception:
            pass

    # If vehicle name is missing, try to derive from selected file name passed via env
    if not meta.get("VEHICLE NAME"):
        sel = os.environ.get(SELECTED_FILE_ENV, "")
        if sel:
            name = os.path.splitext(os.path.basename(sel))[0]
            if name.lower().startswith("final "):
                name = name[6:].lstrip()
            if name:
                meta["VEHICLE NAME"] = name

    return meta


def meta_value(meta: Dict[str, str], key: str) -> str:
    """Fetch a value from meta with a few common key variants."""
    if key in meta:
        return meta[key]
    alt = meta.get(key.title())
    if alt:
        return alt
    alt = meta.get(key.lower())
    if alt:
        return alt
    return ""


def load_output_config(tests_folder: str) -> Dict[int, Dict[str, str]]:
    """Read file_name.json to map result/summary/graph filenames for each test."""
    config_path = os.path.join(tests_folder, FILE_MAP_NAME)
    config_data: Dict[str, Dict[str, str]] = {}
    if os.path.exists(config_path):
        try:
            with open(config_path, "r") as f:
                raw = json.load(f)
            config_data = {k.lower(): v for k, v in raw.items() if isinstance(v, dict)}
        except Exception:
            config_data = {}

    output_by_row: Dict[int, Dict[str, str]] = {}
    for row, script_name in SCRIPT_BY_ROW.items():
        test_name = TEST_CASES[row].lower()
        entry = config_data.get(test_name, {})
        defaults = _default_output_names(script_name)
        output_by_row[row] = {
            "result": entry.get("result", defaults["result"]),
            "summary": entry.get("summary", defaults["summary"]),
            "graph": entry.get("graph", defaults["graph"]),
        }
    return output_by_row


def shade_cell(cell, color_hex: str):
    # w:val is required for valid shading; use "clear" with a fill color.
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    )


def set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    """Apply cell padding (twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for attr, val in (("w:top", top), ("w:left", left), ("w:bottom", bottom), ("w:right", right)):
        if val is None:
            continue
        node = tcMar.find(qn(attr))
        if node is None:
            node = OxmlElement(attr)
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, size=12, color: Tuple[int, int, int] = (0, 0, 0)):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*color)


def add_summary_page(doc: Document, meta: Dict[str, str]):
    vehicle_name = meta.get("VEHICLE NAME", "N/A")
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hrun = heading.add_run(f"VEHICLE NAME : {vehicle_name}")
    hrun.font.size = Pt(16)
    hrun.bold = True

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = False
    try:
        table.columns[0].width = Inches(3.6)
        table.columns[1].width = Inches(2.2)
    except Exception:
        pass
    rows = [
        ("BMS HW VERSION", "#1F4EB2"),
        ("BMS FIRMWARE", "#1F4EB2"),
        ("BMS CONFIG ID", "#1F4EB2"),
        ("BMS GITSHA", "#1F4EB2"),
        ("BMS MANIFEST", "#1F4EB2"),
        ("STARK FIRMWARE", "#00A9E0"),
        ("STARK CONFIG", "#00A9E0"),
        ("XAVIER FIRMWARE", "#109E7D"),
        ("DISTANCE COVERED", "#109E7D"),
    ]
    for label, color in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], color.lstrip("#"))
        set_cell_margins(cells[0])
        set_cell_text(cells[0], label, bold=True, size=11, color=(255, 255, 255))
        shade_cell(cells[1], "F2F2F2")
        set_cell_margins(cells[1])
        set_cell_text(cells[1], meta_value(meta, label), bold=False, size=11, color=(0, 0, 0))

    # Reset section
    reset_table = doc.add_table(rows=0, cols=3)
    reset_table.style = "Table Grid"
    reset_table.allow_autofit = False
    try:
        reset_table.columns[0].width = Inches(3.0)
        reset_table.columns[1].width = Inches(1.3)
        reset_table.columns[2].width = Inches(1.1)
    except Exception:
        pass

    def reset_row(title, count_key, result_key):
        cells = reset_table.add_row().cells
        count_val = meta_value(meta, count_key)
        result_val = meta_value(meta, result_key).upper()
        is_pass = "PASS" in result_val
        is_fail = "FAIL" in result_val
        base_color = "28A745" if is_pass else ("FF0000" if is_fail else "CCCCCC")
        text_color = (255, 255, 255) if (is_pass or is_fail) else (0, 0, 0)

        for c in cells:
            set_cell_margins(c)
            shade_cell(c, base_color)

        set_cell_text(cells[0], title, bold=True, size=11, color=text_color)
        set_cell_text(cells[1], count_val or "Count : 0", bold=True, size=11, color=text_color)
        set_cell_text(cells[2], result_val or "N/A", bold=True, size=11, color=text_color)

    reset_row("VCU unexpected Reset", "VCU Reset Count", "VCU Reset Result")
    reset_row("MARVEL BMS unexpected Reset", "BMS Reset Count", "BMS Reset Result")


def add_ascii_block(doc: Document, lines: List[str]):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, "1F1F1F")
    p = cell.paragraphs[0]
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)


def add_result_line(doc: Document, result: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    color = (0, 122, 60) if str(result).upper() == "PASS" else (204, 0, 0)
    run = para.add_run(f"RESULT : {result or 'N/A'}")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(*color)

def set_page_borders(doc: Document):
    """Apply a thin black border around each page."""
    for section in doc.sections:
        pg_borders = section._sectPr.find(qn("w:pgBorders"))
        if pg_borders is None:
            pg_borders = OxmlElement("w:pgBorders")
            section._sectPr.append(pg_borders)
        pg_borders.set(qn("w:offsetFrom"), "page")
        attrs = {
            qn("w:val"): "single",
            qn("w:sz"): "24",  # thicker border
            qn("w:space"): "12",  # twips
            qn("w:color"): "000000",
        }
        for edge in ("top", "bottom", "left", "right"):
            node = pg_borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                pg_borders.append(node)
            for k, v in attrs.items():
                node.set(k, v)


def _available_width(doc: Document, pad_in: float = 0.1) -> Emu:
    """Compute usable page width (minus margins + a small pad)."""
    section = doc.sections[-1]
    width = section.page_width - section.left_margin - section.right_margin
    pad = Inches(pad_in)
    if width > 2 * pad:
        width -= 2 * pad
    return width


def add_scaled_picture(doc: Document, path: str, max_w_in: float = MAX_GRAPH_WIDTH_IN, max_h_in: float = MAX_GRAPH_HEIGHT_IN):
    """Insert an image scaled to fit the available page width/height while preserving aspect ratio."""
    avail_w = _available_width(doc)
    max_w = min(avail_w, Inches(max_w_in))
    max_h = Inches(max_h_in)
    width_emu = None
    height_emu = None
    try:
        img = Image.from_file(path)
        if img.width and img.height:
            scale = min(max_w / img.width, max_h / img.height, 1.0)
            # If the image is small, allow gentle upscale up to available width
            if img.width * scale < max_w:
                scale = min(max_w / img.width, max_h / img.height)
            width_emu = Emu(int(img.width * scale))
            height_emu = Emu(int(img.height * scale))
    except Exception:
        width_emu = None
        height_emu = None

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run()

    try:
        run.add_picture(path, width=width_emu or max_w, height=height_emu or None)
    except Exception:
        run.add_picture(path, width=max_w)


def _trim_whitespace(path: str) -> str:
    """
    Remove whitespace margins from an image to avoid large blank areas in DOCX.
    Returns path to cropped temp file, or original path on failure.
    """
    try:
        from PIL import Image, ImageChops

        with Image.open(path) as im:
            im = im.convert("RGB")
            bg = Image.new("RGB", im.size, (255, 255, 255))
            diff = ImageChops.difference(im, bg)
            bbox = diff.getbbox()

            # Stronger threshold-based bbox (treat near-white as background)
            if not bbox:
                gray = im.convert("L")
                mask = gray.point(lambda p: 0 if p > 250 else 255)
                bbox = mask.getbbox()

            if not bbox:
                return path

            cropped = im.crop(bbox)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(path)[1] or ".png")
            cropped.save(tmp.name)
            return tmp.name
    except Exception:
        return path


def add_test_page(doc: Document, idx: int, case_name: str, paths: Dict[str, str], tests_folder: str):
    folder = os.path.splitext(SCRIPT_BY_ROW[idx])[0]
    case_folder = os.path.join(tests_folder, folder)
    summary_path = os.path.join(case_folder, paths.get("summary", ""))
    result_path = os.path.join(case_folder, paths.get("result", ""))
    graph_path = os.path.join(case_folder, paths.get("graph", ""))

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.space_after = Pt(2)
    trun = title.add_run(f"TEST CASE {idx + 1} : {case_name.upper()}")
    trun.font.name = "Calibri"
    trun.font.size = Pt(12)
    trun.bold = True
    trun.font.color.rgb = RGBColor(192, 0, 0)  # red

    # Graph first (trim excess whitespace for better layout)
    if os.path.exists(graph_path):
        try:
            trimmed = _trim_whitespace(graph_path)
            add_scaled_picture(doc, trimmed)
        except Exception:
            doc.add_paragraph(f"[Graph could not be inserted: {graph_path}]")
    else:
        doc.add_paragraph(f"[Graph missing: {graph_path}]")

    # Summary block (optional, after image to keep title+graph together)
    skip_summary = case_name.lower() in SKIP_SUMMARY_FOR
    if not skip_summary:
        summary_lines: List[str] = ["Summary not found."]
        if os.path.exists(summary_path):
            try:
                with open(summary_path, "r") as f:
                    data = json.load(f)
                target = data.get("Summary_Table") if isinstance(data, dict) else None
                if target is None:
                    target = data
                text = json.dumps(target, indent=2, ensure_ascii=False)
                summary_lines = text.splitlines() or ["Summary not found."]
            except Exception:
                pass
        add_ascii_block(doc, summary_lines)

    # Result
    result_val = "N/A"
    if os.path.exists(result_path):
        try:
            with open(result_path, "r") as f:
                data = json.load(f)
            result_val = str(data.get("Result", "N/A"))
        except Exception:
            pass
    add_result_line(doc, result_val)


def build_doc(out_path: str, meta_csv: str, tests_folder: str):
    doc = Document()
    set_page_borders(doc)
    meta = load_meta(meta_csv)
    add_summary_page(doc, meta)
    doc.add_page_break()

    output_config = load_output_config(tests_folder)
    for idx, case_name in enumerate(TEST_CASES):
        add_test_page(doc, idx, case_name, output_config.get(idx, {}), tests_folder)
        if idx != len(TEST_CASES) - 1:
            doc.add_page_break()

    doc.save(out_path)
    print(f"Tracker DOCX saved to: {out_path}")


if __name__ == "__main__":
    # Usage: python Generate_Tracker.py [output_docx] [tests_folder] [meta_csv]
    out_path = sys.argv[1] if len(sys.argv) > 1 else OUTPUT_DOCX
    tests_folder = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_TESTS_FOLDER
    meta_csv = sys.argv[3] if len(sys.argv) > 3 else META_CSV

    build_doc(out_path, meta_csv, tests_folder)
